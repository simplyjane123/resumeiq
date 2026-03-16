"""
AI Resume Tailor — Public Edition
Powered by Claude claude-opus-4-6 with adaptive thinking

Upload your own resume, paste a job description, and get a perfectly tailored resume
formatted to the professional reverse-chronological standard.

Usage:
    streamlit run ai_resume_tailor_app.py

Requirements:
    pip install anthropic streamlit python-docx pypdf openai requests
"""

import json
import os
import re
import time
import traceback
import uuid
from datetime import datetime, timezone
from io import BytesIO

import anthropic
import requests
import streamlit as st

# ── Optional providers ─────────────────────────────────────────────────────────
try:
    import openai as _openai_lib
    _OPENAI_AVAILABLE = True
except ImportError:
    _OPENAI_AVAILABLE = False

# ── PDF parsing ────────────────────────────────────────────────────────────────
try:
    import pypdf
    _PYPDF_AVAILABLE = True
except ImportError:
    try:
        import PyPDF2 as pypdf       # legacy fallback
        _PYPDF_AVAILABLE = True
    except ImportError:
        _PYPDF_AVAILABLE = False

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ──────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ──────────────────────────────────────────────────────────────────────────────

TNR = "Calibri"

# ──────────────────────────────────────────────────────────────────────────────
# RESUME UPLOAD & PARSING
# ──────────────────────────────────────────────────────────────────────────────

def parse_resume_upload(uploaded_file) -> str:
    """Extract plain text from an uploaded .docx, .pdf, or .txt resume file."""
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()

    if name.endswith(".docx"):
        doc = Document(BytesIO(raw))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Also capture text inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in lines:
                        lines.append(text)
        return "\n".join(lines)

    elif name.endswith(".pdf"):
        if not _PYPDF_AVAILABLE:
            st.error(
                "PDF support requires the `pypdf` package. "
                "Run: pip install pypdf"
            )
            return ""
        reader = pypdf.PdfReader(BytesIO(raw))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n".join(pages)

    elif name.endswith(".txt") or name.endswith(".md"):
        return raw.decode("utf-8", errors="replace")

    else:
        st.error("Unsupported file type. Please upload a .docx, .pdf, or .txt file.")
        return ""


# ──────────────────────────────────────────────────────────────────────────────
# SYSTEM PROMPTS
# ──────────────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """\
You are a senior hiring manager and expert resume strategist with deep experience tailoring
resumes for competitive roles across all industries.
You apply a value-driven, metrics-first method to transform base resumes into targeted,
ATS-optimised documents.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ABSOLUTE RULES — READ FIRST, APPLY ALWAYS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
❌ NEVER invent, fabricate, or hallucinate:
   - Companies, roles, or employers not in the base resume
   - Projects, products, or initiatives not in the base resume
   - Numbers, metrics, or percentages not in the base resume
   - Skills, tools, certifications, or qualifications not in the base resume
   - Outcomes, results, or impact not in the base resume

✅ You MAY only:
   - Reframe existing experience using the JD's vocabulary
   - Reorder bullets to lead with the most relevant content
   - Tighten or sharpen language for clarity and impact
   - Surface relevant details that were buried or understated
   - Use JD keywords where the underlying experience genuinely supports them

INTEGRITY CHECK (Phase 5): Before finalising every bullet, ask:
  "Is every fact in this sentence explicitly supported by the base resume?"
  If the answer is NO — rewrite or remove it. Never assume or infer new facts.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PHASE 1 — ROLE IDENTIFICATION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Classify the JD and extract: core competencies, required/preferred skills, recurring
keywords, and implicit evaluation criteria (delivery ownership, commercial growth,
stakeholder influence, technical depth, etc.).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PHASE 2 — GAP MAPPING (facts only)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Compare the base resume against the JD:
  - Translation gaps: experience exists but language differs — reframe using JD terms
  - Exposure gaps: partial or adjacent involvement — surface it honestly
  - Real gaps: the experience genuinely does not exist — do NOT fill these in

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PHASE 3 — STRATEGIC REPOSITIONING
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Emphasise experience that directly maps to JD requirements.
Reorder and rewrite bullets to prioritise relevance.
Use JD vocabulary where the underlying experience genuinely supports it.
Ensure bullets are achievement-focused and use only real, sourced numbers.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PHASE 4 — PROFESSIONAL SUMMARY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Rewrite the summary to:
  - Signal trajectory and intent using only existing experience
  - Answer "so what?" in the very first sentence (6-second rule)
  - Use only claims directly supported by the base resume

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PHASE 5 — INTEGRITY REVIEW
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Read every bullet you wrote. For each one, verify:
  ✅ Every company/employer named is in the base resume
  ✅ Every number (%, $, headcount, timeline) is in the base resume
  ✅ Every project or initiative named is in the base resume
  ✅ Every skill or tool claimed is in the base resume
  ✅ The framing is stronger — but the underlying fact is unchanged
  ❌ If any of the above checks fail — rewrite or delete the bullet

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STYLE RULES (apply throughout)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. Write every bullet as a plain, complete sentence — no bold lead-in, no colon header.
   Lead with the result or metric where possible: "Increased X by Y% by doing Z."
   All metrics and results must come from the base resume — never invented.
2. Title line uses the JD's exact role language (2–4 descriptors separated by  •  ).
3. Key Skills uses the JD's exact keywords — only where the resume genuinely supports them.
4. Never write "responsible for" or "assisted with". Use: drove, scaled, owned,
   architected, closed, shipped, secured, optimised, delivered.
5. Metrics front and centre — but only real metrics from the base resume.
6. Every bullet answers: "What did this person actually deliver for the business?"

DOCUMENT FORMAT:
- Font: Calibri throughout. Name: 20pt bold centred. Body text: 11pt.
- Section headers: ALL CAPS bold 11pt with bottom border.
- Page margins: 0.5 inch left/right, 0.625 inch top, 1.0 inch bottom.
- Section order: Name → Contact → EXECUTIVE SUMMARY → KEY SKILLS
  → PROFESSIONAL EXPERIENCE → EARLIER ROLES (if present)
  → EDUCATION & PROFESSIONAL QUALIFICATIONS → COURSES & CERTIFICATIONS
  → OTHER INFORMATION.
- Key Skills: exactly 6 bold keywords in a bullet list — use JD vocabulary.
- All bullets: • flush left, wrapped text hanging-indented to align with first word.
- Each role: company name (bold) on line 1, Job Title | Period (bold) on line 2,
  responsibility bullets, then "Achievements:" bold sub-label + achievement bullets.

OUTPUT FORMAT:
Return ONLY a valid JSON object — no markdown fences, no explanation, no preamble.
The JSON must exactly match the schema described in the user prompt.
"""

COVER_LETTER_SYSTEM_PROMPT = """\
You are an expert cover letter writer. You apply value-first communication principles
that put the employer's needs before the candidate's wants.

CORE PRINCIPLES:
1. Lead with THEIR problem or opportunity — not what the candidate wants. The opening
   must answer "what does this organisation gain?" before introducing the candidate.
2. Balance "you/your" vs "I/my" — lean toward the reader's perspective throughout.
3. Substantiate every claim with evidence, metrics, or specifics.
4. Every paragraph must pass the "could they forward this?" test.
5. Quantify wherever possible — numbers make abstract claims concrete and credible.

COVER LETTER STRUCTURE:
Para 1 HOOK: Open with the organisation's mission, challenge, or goal. Connect the
          candidate's specific value to it. Name the role. Do NOT open with "I am applying for..."
Para 2 BACKGROUND NARRATIVE: Career arc and why the candidate is uniquely positioned.
Para 3 ACHIEVEMENT EVIDENCE: 2–3 specific, metric-backed achievements directly mapped to the JD.
          Draw from base resume only. Make the hiring manager picture the impact.
Para 4 WHY THIS COMPANY + LEARNING: Genuine, specific interest in THIS organisation's mission
          or challenge, combined with relevant recent learning that addresses JD requirements.
Para 5 CLOSING: Confident, specific value statement — what the candidate will deliver for them.
          No weak phrases: never "I hope", "I would be grateful", "I am passionate about".

STYLE RULES:
- Professional but warm tone. Calibri 11pt in the final document.
- Specific over generic — every sentence must earn its place.
- Active verbs: owned, drove, shipped, secured, built, architected, translated, scaled.
- STRICT ONE PAGE: write exactly 4–5 paragraphs, each 3–4 sentences maximum.

INTEGRITY RULES:
- Every achievement, metric, and claim must be sourced from the BASE RESUME.
- Do NOT invent projects, outcomes, companies, numbers, or skills.

OUTPUT FORMAT:
Return ONLY a valid JSON object, no markdown fences, no explanation:
{
  "subject_line": "Application for [Role] – [Company]",
  "paragraphs": ["<para1>", "<para2>", ...],
  "signoff_name": "<candidate full name from resume>"
}
"""


# ──────────────────────────────────────────────────────────────────────────────
# HELPER: build dynamic JSON schema for the user prompt
# ──────────────────────────────────────────────────────────────────────────────

def _resume_json_schema() -> str:
    """Return the JSON schema block that Claude must populate from the uploaded resume."""
    return """{
  "name": "<candidate full name from resume>",
  "contact": "<contact line from resume — phone | email | LinkedIn | location>",
  "title_line": "<2–4 role descriptors using the JD's exact language, separated by  •  >",
  "summary": "<3–4 sentences. Open with a strong value statement that maps directly to this role. Embed 3–5 JD keywords naturally. Use only facts from the base resume.>",
  "key_skills": [
    "<JD keyword — backed by resume>",
    "<JD keyword — backed by resume>",
    "<JD keyword — backed by resume>",
    "<JD keyword — backed by resume>",
    "<JD keyword — backed by resume>",
    "<JD keyword — backed by resume>"
  ],
  "experience": [
    {
      "employer": "<company name from resume>",
      "employer_desc": "<brief company description if mentioned in resume, else empty string>",
      "period": "<employment period from resume>",
      "title": "<job title from resume>",
      "responsibilities": [
        "<Plain sentence describing what was done and its impact>",
        "<Plain sentence describing what was done and its impact>",
        "<Plain sentence describing what was done and its impact>"
      ],
      "achievements": [
        "<Plain sentence leading with a metric or result — e.g. 'Increased X by Y% by doing Z'>",
        "<Plain sentence leading with a metric or result>"
      ]
    }
  ],
  "earlier_roles": [
    {"employer": "<company>", "title": "<title>", "period": "<period>"}
  ],
  "education": [
    {"degree": "<degree name>", "institution": "<school/university>", "year": "<year>"}
  ],
  "courses_certifications": [
    {"name": "<course or certification name>", "body": "<awarding body>", "year": "<year>"}
  ],
  "other_info": {
    "technical_skills": "<comma or bullet separated list of tools/software>",
    "languages": "<languages spoken>",
    "interests": "<personal interests>"
  }
}"""


# ──────────────────────────────────────────────────────────────────────────────
# JSON PARSING UTILITIES
# ──────────────────────────────────────────────────────────────────────────────

def _strip_fences(text: str) -> str:
    text = re.sub(r"^```(?:json)?\s*", "", text.strip(), flags=re.MULTILINE)
    return re.sub(r"\s*```$", "", text, flags=re.MULTILINE)


def _parse_json_response(text: str) -> dict:
    """Robustly extract and parse the first JSON object from an LLM response.

    Handles:
    - Markdown code fences (```json ... ```)
    - Preamble / explanation text before or after the JSON
    - Unescaped newlines / stray characters inside string values (best-effort)
    - Truncated responses (adds missing closing braces)
    """
    text = _strip_fences(text).strip()

    # Fast path — try direct parse first
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Find the outermost { … } using bracket counting that respects strings
    start = text.find("{")
    if start == -1:
        # Try array (for score_jobs_with_claude which returns [])
        start = text.find("[")
        if start == -1:
            raise json.JSONDecodeError("No JSON object found in response", text, 0)

    opener = text[start]
    closer = "}" if opener == "{" else "]"

    depth = 0
    in_string = False
    escape_next = False
    end = -1

    for i, ch in enumerate(text[start:], start):
        if escape_next:
            escape_next = False
            continue
        if ch == "\\" and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if not in_string:
            if ch == opener:
                depth += 1
            elif ch == closer:
                depth -= 1
                if depth == 0:
                    end = i
                    break

    if end == -1:
        # Response was truncated — close any open braces and try anyway
        fragment = text[start:]
        open_count = fragment.count(opener) - fragment.count(closer)
        fragment += closer * max(open_count, 1)
        return json.loads(fragment)

    json_str = text[start : end + 1]

    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        # Last-resort: replace literal newlines inside strings with \\n
        json_str_fixed = re.sub(
            r'(?<!\\)"((?:[^"\\]|\\.)*)"',
            lambda m: '"' + m.group(1).replace("\n", "\\n").replace("\r", "") + '"',
            json_str,
        )
        return json.loads(json_str_fixed)


# ──────────────────────────────────────────────────────────────────────────────
# CLAUDE API — RESUME TAILORING
# ──────────────────────────────────────────────────────────────────────────────

def tailor_with_claude(
    base_resume: str,
    job_description: str,
    company_name: str,
    role_name: str,
    gaps: list | None = None,
    missing_keywords: list | None = None,
) -> dict:
    """Call Claude claude-opus-4-6 to tailor the uploaded resume and return parsed JSON."""
    client = anthropic.Anthropic()

    gap_block = ""
    if gaps or missing_keywords:
        gap_lines = []
        if gaps:
            gap_lines.append("Identified gaps to address:\n" + "\n".join(f"  - {g}" for g in gaps))
        if missing_keywords:
            gap_lines.append(
                "Missing ATS keywords to weave in (where truthful):\n  " + ", ".join(missing_keywords)
            )
        gap_block = (
            "\n\nGAP-CLOSING INSTRUCTIONS (from job match analysis):\n"
            "The following gaps and missing keywords were identified. "
            "Strategically address each gap by reframing or highlighting relevant experience. "
            "Incorporate missing keywords naturally where accurate.\n"
            + "\n".join(gap_lines)
        )

    schema = _resume_json_schema()

    user_prompt = f"""\
BASE RESUME (uploaded by user):
{base_resume}

JOB TO TAILOR FOR:
Company: {company_name}
Role: {role_name}
Job Description:
{job_description}{gap_block}

Return a JSON object with this exact schema. Populate ALL fields from the uploaded
resume above. Include ALL substantive roles (reverse chronological). Do NOT use
placeholder text — every field must contain real content from the base resume.

Rules for the "experience" array:
- Include each substantive role as a separate entry (do not merge roles).
- Include 3–6 responsibility bullets and 2–4 achievement bullets per role.
- If a role has no listed achievements, set "achievements" to [].
- Put very early or brief roles (e.g., <1 year, 10+ years ago) in "earlier_roles" instead.

Rules for "earlier_roles":
- Brief one-line entries only: employer, title, period.
- Include only if they appear in the base resume.

{schema}

IMPORTANT — INTEGRITY RULES (strictly enforced):
- Every employer, project, metric, number, and skill in the output must exist verbatim
  or be directly inferable from the BASE RESUME above. Do NOT invent anything.
- Do NOT add new projects, initiatives, tools, qualifications, or companies.
- Do NOT change any number (%, $, headcount, timeline). Use exact figures only.
- Only reframe language and reorder emphasis — the underlying facts must stay intact.
- If the base resume does not support a JD requirement, leave it out entirely.
- Return ONLY the JSON object — no markdown, no explanation.
"""

    _thinking_configs = [
        {"type": "adaptive"},
        {"type": "enabled", "budget_tokens": 8000},
        {"type": "enabled", "budget_tokens": 4000},
    ]

    for _attempt in range(3):
        with st.spinner(
            f"Claude is tailoring your resume"
            f"{'…' if _attempt == 0 else f' (retry {_attempt}/2)…'}"
        ):
            try:
                response = client.messages.create(
                    model="claude-opus-4-6",
                    max_tokens=16000,
                    thinking=_thinking_configs[_attempt],
                    system=SYSTEM_PROMPT,
                    messages=[{"role": "user", "content": user_prompt}],
                )
            except anthropic.APIStatusError as _e:
                if _e.status_code == 529 and _attempt < 2:
                    _wait = [15, 30][_attempt]
                    st.warning(f"Anthropic servers busy — retrying in {_wait}s…")
                    time.sleep(_wait)
                    continue
                raise

        text = next((b.text for b in response.content if b.type == "text"), "")
        if not text.strip():
            if _attempt < 2:
                st.warning(f"Claude returned an empty response (attempt {_attempt + 1}/3) — retrying…")
                continue
            raise ValueError("Claude returned an empty response after all retries. Please try again.")

        try:
            return _parse_json_response(text)
        except json.JSONDecodeError as _je:
            if _attempt < 2:
                st.warning(f"Claude returned malformed JSON (attempt {_attempt + 1}/3) — retrying…")
                continue
            preview = text[:500] if text else "(empty)"
            raise json.JSONDecodeError(
                f"Claude returned invalid JSON after 3 attempts. Response preview: {preview}",
                text, _je.pos,
            ) from _je

    raise RuntimeError("tailor_with_claude: exceeded retry limit")


# ──────────────────────────────────────────────────────────────────────────────
# CLAUDE API — COVER LETTER
# ──────────────────────────────────────────────────────────────────────────────

def generate_cover_letter(
    base_resume: str,
    job_description: str,
    company_name: str,
    role_name: str,
    gaps: list | None = None,
    missing_keywords: list | None = None,
) -> dict:
    """Call Claude to generate a value-driven cover letter and return parsed JSON."""
    client = anthropic.Anthropic()

    gap_block = ""
    if gaps or missing_keywords:
        parts = []
        if gaps:
            parts.append("Gaps to address:\n" + "\n".join(f"  - {g}" for g in gaps))
        if missing_keywords:
            parts.append("Missing keywords to incorporate (where truthful):\n  " + ", ".join(missing_keywords))
        gap_block = (
            "\n\nGAP-CLOSING INSTRUCTIONS:\n"
            "Address these gaps by reframing the candidate's experience to speak to these dimensions, "
            "and use the missing keywords naturally where accurate.\n"
            + "\n".join(parts)
        )

    user_prompt = f"""\
BASE RESUME:
{base_resume}

JOB TO WRITE COVER LETTER FOR:
Company: {company_name}
Role: {role_name}
Job Description:
{job_description}{gap_block}

Write a cover letter following your system instructions.

INTEGRITY RULES:
- Every achievement, metric, and claim must be directly sourced from the BASE RESUME.
- Do NOT invent projects, outcomes, companies, numbers, or skills.
- Use exact figures — do not round, estimate, or upgrade them.
- Return ONLY the JSON object.
"""

    _retry_delays = [10, 20, 40]
    for _attempt in range(len(_retry_delays) + 1):
        try:
            response = client.messages.create(
                model="claude-opus-4-6",
                max_tokens=3000,
                thinking={"type": "enabled", "budget_tokens": 2000},
                system=COVER_LETTER_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_prompt}],
            )
            break
        except anthropic.APIStatusError as _e:
            if _e.status_code == 529 and _attempt < len(_retry_delays):
                st.warning(f"Anthropic servers busy — retrying cover letter in {_retry_delays[_attempt]}s…")
                time.sleep(_retry_delays[_attempt])
            else:
                raise
        except (anthropic.APIConnectionError, Exception) as _e:
            if _attempt < len(_retry_delays):
                st.warning(f"Connection error on cover letter — retrying in {_retry_delays[_attempt]}s… ({_e})")
                time.sleep(_retry_delays[_attempt])
            else:
                raise

    text = next((b.text for b in response.content if b.type == "text"), "")
    return _parse_json_response(text)


# ──────────────────────────────────────────────────────────────────────────────
# OPENAI — SECOND-PASS REFINEMENT
# ──────────────────────────────────────────────────────────────────────────────

def refine_with_openai(
    base_resume: str,
    tailored: dict,
    job_description: str,
    gaps: list | None = None,
    missing_keywords: list | None = None,
) -> dict:
    """Second-tier refinement: GPT-4o reviews and strengthens the Claude-tailored JSON."""
    client = _openai_lib.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    gap_note = ""
    if gaps or missing_keywords:
        parts = []
        if gaps:
            parts.append("Gaps still to address: " + "; ".join(gaps))
        if missing_keywords:
            parts.append("Missing ATS keywords to weave in (where truthful): " + ", ".join(missing_keywords))
        gap_note = "\n\n" + "\n".join(parts)

    prompt = f"""You are an expert resume editor performing a second-tier refinement pass.

Claude has already tailored this resume JSON for a specific job. Your job is to strengthen it further.

JOB DESCRIPTION (first 2000 chars):
{job_description[:2000]}
{gap_note}

BASE RESUME (for fact-checking):
{base_resume[:2000]}

CLAUDE'S TAILORED RESUME JSON:
{json.dumps(tailored, ensure_ascii=False, indent=2)}

YOUR TASK — improve the following without changing the JSON schema or fabricating facts:
1. **Bullets**: Each bullet is a plain sentence — no bold lead-in, no colon header.
   Lead with the result or metric where possible: "Increased X by Y% by doing Z."
   Sharpen action verbs and tighten language.
2. **Summary**: Make the opening sentence even more punchy and role-specific.
   Ensure it answers "so what?" in under 10 words.
3. **Title line**: Confirm it uses the JD's exact language. Tighten if needed.
4. **Key Skills**: Verify all 6 items are JD-keyword-aligned. Replace weak generic terms.
5. **Gap keywords**: Ensure missing ATS keywords are naturally present where accurate.

INTEGRITY RULES:
- Every employer, project, metric, number, and skill must exist in the base resume.
- Do NOT invent, add, or imply anything not explicitly in the base resume.
- Do NOT change any number (%, $, headcount, timeline) — use exact figures only.
- Do NOT add new projects, tools, qualifications, or companies.
- Return the EXACT same JSON schema — do not add or remove keys.
- Do not modify: earlier_roles, education, courses_certifications, other_info, contact, name.
- Return ONLY the JSON object, no markdown fences, no explanation."""

    response = client.chat.completions.create(
        model="gpt-4o",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )
    return _parse_json_response(response.choices[0].message.content)


# ──────────────────────────────────────────────────────────────────────────────
# JOB MATCHER — MyCareerFuture API + Scoring
# ──────────────────────────────────────────────────────────────────────────────

MCF_API_BASE = "https://api.mycareersfuture.gov.sg/v2"


def fetch_mcf_jobs(keywords: str, limit: int = 20) -> list:
    """Fetch jobs from the MyCareerFuture public API (Singapore)."""
    try:
        resp = requests.get(
            f"{MCF_API_BASE}/jobs",
            params={
                "search": keywords,
                "limit": limit,
                "sortBy": "new_posting_date",
                "descending": "true",
            },
            headers={"Accept": "application/json", "User-Agent": "AIResumeTailorApp/1.0"},
            timeout=15,
        )
        resp.raise_for_status()
        data = resp.json()
        jobs = []
        for item in data.get("results", []):
            desc_raw = (item.get("description") or {}).get("content", "")
            desc_clean = re.sub(r"<[^>]+>", " ", desc_raw)
            desc_clean = re.sub(r"\s+", " ", desc_clean).strip()
            meta = item.get("metadata") or {}
            posted_raw = (
                meta.get("newPostingDate")
                or meta.get("createdAt")
                or item.get("createdAt", "")
            )
            sal = item.get("salary") or {}
            jobs.append({
                "title": item.get("title", "Untitled"),
                "company": (item.get("postedCompany") or {}).get("name", "Unknown"),
                "description": desc_clean[:2000],
                "posted_raw": posted_raw,
                "uuid": item.get("uuid", ""),
                "salary_min": sal.get("minimum"),
                "salary_max": sal.get("maximum"),
            })
        return jobs
    except requests.exceptions.RequestException as exc:
        st.error(f"Could not reach MyCareerFuture API: {exc}")
        return []
    except Exception as exc:
        st.error(f"Error fetching jobs: {exc}")
        return []


def _parse_posted_date(posted_raw: str):
    if not posted_raw:
        return None
    try:
        return datetime.fromisoformat(posted_raw.replace("Z", "+00:00"))
    except Exception:
        return None


def _jd_match_prompt(base_resume: str, jd_text: str) -> str:
    return f"""\
CANDIDATE RESUME:
{base_resume}

JOB DESCRIPTION:
{jd_text[:3000]}

Analyse how well this candidate's resume matches the job description.

Return a JSON object:
{{
  "score": <integer 0-100>,
  "fit": "<one sentence summary of overall fit>",
  "match": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "gap": ["<gap or missing requirement 1>", "<gap 2>"],
  "recommendation": "<'Apply' or 'Skip'>",
  "recommendation_reason": "<1-2 sentence plain-English explanation>",
  "missing_keywords": ["<ATS keyword missing from resume>", "<keyword 2>", "<keyword 3>"]
}}

Scoring: 85-100 near-perfect, 70-84 strong, 50-69 moderate, below 50 weak.
Recommend 'Apply' if score >= 65, otherwise 'Skip'.
Return ONLY the JSON object, no markdown fences."""


def score_pasted_jd(base_resume: str, jd_text: str) -> dict:
    """Score a single pasted JD against the uploaded resume using Claude Haiku."""
    client = anthropic.Anthropic()
    _retry_delays = [10, 20, 40]
    for _attempt in range(len(_retry_delays) + 1):
        try:
            response = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=1000,
                messages=[{"role": "user", "content": _jd_match_prompt(base_resume, jd_text)}],
            )
            return _parse_json_response(response.content[0].text)
        except anthropic.APIStatusError as _e:
            if _e.status_code == 529 and _attempt < len(_retry_delays):
                time.sleep(_retry_delays[_attempt])
            else:
                raise
        except (anthropic.APIConnectionError, Exception) as _e:
            if _attempt < len(_retry_delays):
                time.sleep(_retry_delays[_attempt])
            else:
                raise


def score_pasted_jd_openai(base_resume: str, jd_text: str) -> dict:
    """Score a single pasted JD against the uploaded resume using OpenAI GPT-4o mini."""
    client = _openai_lib.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        max_tokens=1000,
        messages=[{"role": "user", "content": _jd_match_prompt(base_resume, jd_text)}],
    )
    return _parse_json_response(response.choices[0].message.content)


def score_jobs_with_claude(jobs: list, base_resume: str) -> list:
    """Score a batch of MCF jobs against the uploaded resume using Claude Haiku."""
    if not jobs:
        return jobs
    client = anthropic.Anthropic()
    jobs_text = ""
    for i, job in enumerate(jobs, 1):
        sal = ""
        if job.get("salary_min") or job.get("salary_max"):
            sal = f"  |  SGD {job.get('salary_min', '?')} – {job.get('salary_max', '?')}/mth"
        jobs_text += (
            f"\nJOB {i}: {job['title']} at {job['company']}{sal}\n"
            f"{job['description'][:500]}\n"
        )
    prompt = f"""\
CANDIDATE RESUME:
{base_resume}

JOBS TO SCORE ({len(jobs)} total):
{jobs_text}

Return a JSON array — one object per job, same order:
[
  {{
    "score": <integer 0–100>,
    "fit": "<10-word summary of fit>",
    "match": ["<top reason 1>", "<top reason 2>"],
    "gap": ["<main gap or missing requirement>"]
  }},
  ...
]
Scoring guide: 85–100 near-perfect, 70–84 strong, 50–69 moderate, below 50 weak.
Return ONLY the JSON array, no markdown."""
    _retry_delays = [15, 30, 60]
    for _attempt in range(len(_retry_delays) + 1):
        try:
            response = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=2500,
                messages=[{"role": "user", "content": prompt}],
            )
            break
        except anthropic.APIStatusError as _e:
            if _e.status_code == 529 and _attempt < len(_retry_delays):
                time.sleep(_retry_delays[_attempt])
            else:
                raise
    text = response.content[0].text.strip()
    scores = _parse_json_response(text)
    for i, s in enumerate(scores):
        if i < len(jobs):
            jobs[i]["score"] = int(s.get("score", 0))
            jobs[i]["fit"] = s.get("fit", "")
            jobs[i]["match"] = s.get("match", [])
            jobs[i]["gap"] = s.get("gap", [])
    return jobs


# ──────────────────────────────────────────────────────────────────────────────
# PROFILE ANALYSIS (session-based — no file persistence)
# ──────────────────────────────────────────────────────────────────────────────

def get_profile_data() -> dict:
    if "profile_data" not in st.session_state:
        st.session_state["profile_data"] = {
            "analyzed_jds": [],
            "tailored_roles": [],
            "profile_analysis": {},
        }
    return st.session_state["profile_data"]


def generate_profile_analysis(profile_data: dict, base_resume: str) -> str:
    """Call Claude Haiku to synthesise career insights from accumulated JD data."""
    analyzed = profile_data.get("analyzed_jds", [])
    tailored = profile_data.get("tailored_roles", [])

    jd_summaries = "\n".join(
        f"- Score: {j['score']}% | Rec: {j['recommendation']} | "
        f"Fit: {j.get('fit', '')} | Gaps: {', '.join(j.get('gap', []))} | "
        f"Missing keywords: {', '.join(j.get('missing_keywords', []))}"
        for j in analyzed
    )
    tailored_summaries = "\n".join(
        f"- {r['role']} at {r['company']} ({r['timestamp'][:10]})"
        for r in tailored
    ) or "None yet"

    prompt = f"""Based on the following job search activity, provide a structured career profile analysis.

CANDIDATE BACKGROUND (uploaded resume excerpt):
{base_resume[:1500]}

JOB DESCRIPTIONS ANALYSED ({len(analyzed)} total):
{jd_summaries}

ROLES WHERE RESUME WAS TAILORED:
{tailored_summaries}

Provide a detailed markdown analysis covering exactly these sections:

## 🎯 Best-Fit Role Types
The specific role titles and function types this candidate is strongest for.

## 🏢 Target Industries
Industries showing the best alignment and why.

## ⭐ Competitive Strengths
The 3–5 differentiators that keep surfacing as match reasons.

## 📈 Skills & Gaps to Address
Recurring gaps that would open more doors if closed.

## 🔍 Where to Focus the Job Search
Specific, actionable targeting advice — role levels, company types, channels.

## 💡 Strategic Recommendations
Patterns or trends that should shape the candidate's overall approach.

Be specific and direct. Reference actual patterns from the data. Use bullet points within each section."""

    client = anthropic.Anthropic()
    _retry_delays = [15, 30, 60]
    for _attempt in range(len(_retry_delays) + 1):
        try:
            response = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=2000,
                messages=[{"role": "user", "content": prompt}],
            )
            return response.content[0].text
        except anthropic.APIStatusError as _e:
            if _e.status_code == 529 and _attempt < len(_retry_delays):
                time.sleep(_retry_delays[_attempt])
            else:
                raise


# ──────────────────────────────────────────────────────────────────────────────
# DOCX BUILDER — matching the professional reverse-chronological template
# Calibri throughout | 20pt bold centred name | 11pt body
# Section headers: ALL CAPS bold with bottom border
# Margins: 0.5" L/R · 0.625" top · 1.0" bottom
# ──────────────────────────────────────────────────────────────────────────────

def _set_font(run, size: float, bold: bool = None, italic: bool = None):
    run.font.name = TNR
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    # Keep East-Asia font consistent
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    rf.set(qn("w:eastAsia"), TNR)


def _new_para(doc, style="No Spacing", space_before=0, space_after=0):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style]
    except KeyError:
        p.style = doc.styles["Normal"]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _section_header(doc, title: str, space_before: int = 6):
    """ALL CAPS bold 11pt section heading with a bottom border line (matching template)."""
    p = _new_para(doc, space_before=space_before, space_after=2)
    r = p.add_run(title)
    _set_font(r, 11, bold=True)
    # Add bottom paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_right_tab(p, twips: int = 10972):
    """Right-aligned tab stop at the right margin (~7.62 inches from left edge)."""
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(twips))
    tabs.append(tab)
    pPr.append(tabs)


def _remove_table_borders(tbl):
    """Strip all borders from every cell in a table."""
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                tcBorders.append(el)
            tcPr.append(tcBorders)


def _plain_bullet(doc, text: str, size: float = 11):
    """Plain bullet — no bold, just a sentence."""
    p = _new_para(doc, style="List Paragraph", space_before=1, space_after=1)
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.first_line_indent = Pt(-14)
    r0 = p.add_run("• ")
    _set_font(r0, size, bold=False)
    r1 = p.add_run(text)
    _set_font(r1, size, bold=False)
    return p


def _bold_bullet(doc, lead: str, rest: str, size: float = 11):
    """Bullet with bold lead-in; flush-left with hanging indent for wrapped lines."""
    p = _new_para(doc, style="List Paragraph", space_before=1, space_after=1)
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.first_line_indent = Pt(-14)
    r0 = p.add_run("• ")
    _set_font(r0, size, bold=False)
    if lead:
        r1 = p.add_run(lead)
        _set_font(r1, size, bold=True)
    if rest:
        r2 = p.add_run(rest)
        _set_font(r2, size, bold=False)
    return p


def build_docx(data: dict) -> BytesIO:
    """Assemble a fully-formatted .docx from the tailored JSON data.

    Format matches: Resume Template - Reverse Chronological Resume - 1.docx
    """
    doc = Document()

    # ── Page layout ──────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.625)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(0.5)
    sec.right_margin  = Inches(0.5)

    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)

    # ── Name (centred, 20pt bold) ─────────────────────────────────────────────
    p = _new_para(doc)
    p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get("name", ""))
    _set_font(r, 20, bold=True)

    # ── Contact line (centred, 11pt bold) ─────────────────────────────────────
    p = _new_para(doc, space_before=2)
    p.alignment = 1
    r = p.add_run(data.get("contact", ""))
    _set_font(r, 11, bold=True)

    # ── EXECUTIVE SUMMARY ────────────────────────────────────────────────────
    _section_header(doc, "EXECUTIVE SUMMARY")
    p = _new_para(doc, style="Normal", space_before=2)
    r = p.add_run(data.get("summary", ""))
    _set_font(r, 11)

    # ── KEY SKILLS (2 columns × 3 rows borderless table) ─────────────────────
    _section_header(doc, "KEY SKILLS")
    skills = [s for s in data.get("key_skills", []) if s][:6]
    while len(skills) < 6:
        skills.append("")
    tbl = doc.add_table(rows=3, cols=2)
    _remove_table_borders(tbl)
    for row_idx in range(3):
        for col_idx in range(2):
            skill_text = skills[row_idx + col_idx * 3]
            cell = tbl.cell(row_idx, col_idx)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            if skill_text:
                p = cell.paragraphs[0]
                r0 = p.add_run("• ")
                _set_font(r0, 11, bold=False)
                r1 = p.add_run(skill_text)
                _set_font(r1, 11, bold=True)

    # ── PROFESSIONAL EXPERIENCE ───────────────────────────────────────────────
    _section_header(doc, "PROFESSIONAL EXPERIENCE")

    for exp in data.get("experience", []):
        # Line 1: Company name (bold) + optional description
        p = _new_para(doc, space_before=6)
        r1 = p.add_run(exp.get("employer", ""))
        _set_font(r1, 11, bold=True)
        desc = exp.get("employer_desc", "")
        if desc:
            r2 = p.add_run(f"  ({desc})")
            _set_font(r2, 11)

        # Line 2: Job Title | Period (bold, right-aligned tab for date)
        p = _new_para(doc, space_before=1)
        _add_right_tab(p)
        r1 = p.add_run(exp.get("title", ""))
        _set_font(r1, 11, bold=True)
        r2 = p.add_run(f"\t{exp.get('period', '')}")
        _set_font(r2, 11, bold=True)

        # Responsibility bullets
        for b in exp.get("responsibilities", []):
            text = b if isinstance(b, str) else (b.get("lead", "") + " " + b.get("rest", "")).strip()
            if text:
                _plain_bullet(doc, text)

        # Achievements sub-section
        raw_achievements = exp.get("achievements", [])
        achievements = []
        for a in raw_achievements:
            text = a if isinstance(a, str) else (a.get("lead", "") + " " + a.get("rest", "")).strip()
            if text:
                achievements.append(text)
        if achievements:
            p = _new_para(doc, space_before=3)
            r = p.add_run("Achievements:")
            _set_font(r, 11, bold=True)
            for text in achievements:
                _plain_bullet(doc, text)

    # ── EARLIER ROLES ─────────────────────────────────────────────────────────
    earlier = data.get("earlier_roles", [])
    if earlier:
        _section_header(doc, "EARLIER ROLES")
        for role in earlier:
            p = _new_para(doc, space_before=2)
            _add_right_tab(p)
            r1 = p.add_run(f"{role.get('employer', '')}    ")
            _set_font(r1, 11, bold=True)
            r2 = p.add_run(f"{role.get('title', '')}\t{role.get('period', '')}")
            _set_font(r2, 11)

    # ── EDUCATION & PROFESSIONAL QUALIFICATIONS ───────────────────────────────
    education = data.get("education", [])
    if education:
        _section_header(doc, "EDUCATION & PROFESSIONAL QUALIFICATIONS")
        for edu in education:
            p = _new_para(doc, space_before=3)
            r1 = p.add_run(edu.get("degree", ""))
            _set_font(r1, 11, bold=True)
            institution = edu.get("institution", "")
            year = edu.get("year", "")
            if institution or year:
                sep = f"  {institution}, {year}" if institution and year else f"  {institution}{year}"
                r2 = p.add_run(sep)
                _set_font(r2, 11)

    # ── COURSES & CERTIFICATIONS ──────────────────────────────────────────────
    certs = data.get("courses_certifications", [])
    if certs:
        _section_header(doc, "COURSES & CERTIFICATIONS")
        for cert in certs:
            name = cert.get("name", "")
            body = cert.get("body", "")
            year = cert.get("year", "")
            p = _new_para(doc, space_before=2)
            parts = []
            if name:
                parts.append(name)
            suffix = ""
            if body and year:
                suffix = f"  –  {body} ({year})"
            elif body:
                suffix = f"  –  {body}"
            elif year:
                suffix = f"  ({year})"
            r = p.add_run("".join(parts) + suffix)
            _set_font(r, 11)

    # ── OTHER INFORMATION ─────────────────────────────────────────────────────
    other = data.get("other_info", {})
    if other and any(other.get(k) for k in ("technical_skills", "languages", "interests")):
        _section_header(doc, "OTHER INFORMATION")
        if other.get("technical_skills"):
            p = _new_para(doc, space_before=3)
            r1 = p.add_run("Technical Skills:  ")
            _set_font(r1, 11, bold=True)
            r2 = p.add_run(other["technical_skills"])
            _set_font(r2, 11)
        if other.get("languages"):
            p = _new_para(doc, space_before=2)
            r1 = p.add_run("Languages:  ")
            _set_font(r1, 11, bold=True)
            r2 = p.add_run(other["languages"])
            _set_font(r2, 11)
        if other.get("interests"):
            p = _new_para(doc, space_before=2)
            r1 = p.add_run("Interests:  ")
            _set_font(r1, 11, bold=True)
            r2 = p.add_run(other["interests"])
            _set_font(r2, 11)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────────────────────────────────────
# COVER LETTER DOCX BUILDER
# ──────────────────────────────────────────────────────────────────────────────

def build_cover_letter_docx(data: dict, company_name: str, role_name: str) -> BytesIO:
    """Build a formatted cover letter .docx from the Claude JSON output."""
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1)
    sec.right_margin  = Inches(1)

    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)

    def _cl_para(text, bold=False, italic=False, size=11, space_before=0, space_after=8):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name  = TNR
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.italic = italic
        return p

    # Date
    _cl_para(datetime.now().strftime("%d %B %Y"), size=11, space_after=8)

    # Subject line (bold)
    subject = data.get(
        "subject_line",
        f"Application for {role_name} \u2013 {company_name}"
    )
    _cl_para(subject, bold=True, size=11, space_after=8)

    # Greeting
    _cl_para("Dear Hiring Manager,", size=11, space_after=8)

    # Body paragraphs
    for para in data.get("paragraphs", []):
        _cl_para(para, size=11, space_after=8)

    # Sign-off
    candidate_name = data.get("signoff_name", "")
    _cl_para("Sincerely,", size=11, space_before=12, space_after=4)
    if candidate_name:
        _cl_para(candidate_name, bold=True, size=11, space_after=0)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────────────────────────────────────
# STREAMLIT UI
# ──────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="ResumeIQ",
    page_icon="🧠",
    layout="wide",
)

st.title("🧠 ResumeIQ")
st.markdown("*Upload your resume · Paste a job description · Get a perfectly tailored resume*")

# ── API key check ──────────────────────────────────────────────────────────────
api_key = os.environ.get("ANTHROPIC_API_KEY", "")
if not api_key:
    st.warning(
        "⚠️ `ANTHROPIC_API_KEY` not set. "
        "Add it to your environment or Streamlit secrets before using this tool."
    )

# ── Sidebar — Resume Upload ───────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 📎 Your Resume")
    st.markdown(
        "Upload your resume once. It will be used across all tabs — "
        "tailoring, job matching, and analysis."
    )

    uploaded_file = st.file_uploader(
        "Upload resume",
        type=["docx", "pdf", "txt"],
        label_visibility="collapsed",
        key="resume_uploader",
    )

    if uploaded_file is not None:
        # Only re-parse if a new file was uploaded
        if (
            "resume_filename" not in st.session_state
            or st.session_state.get("resume_filename") != uploaded_file.name
        ):
            with st.spinner("Parsing your resume…"):
                parsed = parse_resume_upload(uploaded_file)
            if parsed:
                st.session_state["user_resume"] = parsed
                st.session_state["resume_filename"] = uploaded_file.name
                # Clear profile data when a new resume is loaded
                st.session_state.pop("profile_data", None)
                st.success(f"✅ Loaded: **{uploaded_file.name}**")
            else:
                st.error("Could not extract text from this file.")
        else:
            st.success(f"✅ Loaded: **{uploaded_file.name}**")

    if "user_resume" in st.session_state:
        with st.expander("Preview extracted text", expanded=False):
            st.text(st.session_state["user_resume"][:1500] + "…")
        if st.button("🗑 Clear resume", use_container_width=True):
            for k in ("user_resume", "resume_filename", "profile_data", "tailor_result"):
                st.session_state.pop(k, None)
            st.rerun()
    else:
        st.info("No resume loaded yet. Upload a .docx, .pdf, or .txt file above.")

    st.markdown("---")
    st.markdown("**Powered by**")
    st.markdown("🧠 Claude claude-opus-4-6 (tailoring)")
    if _OPENAI_AVAILABLE and os.environ.get("OPENAI_API_KEY"):
        st.markdown("✨ GPT-4o (refinement pass)")
    st.markdown("⚡ Claude Haiku (scoring)")


# ── Tabs ──────────────────────────────────────────────────────────────────────
tab1, tab2, tab3 = st.tabs(["🔍 Job Matcher", "✏️ Tailor Resume", "📊 Analysis"])


# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — Job Matcher  (moved first so users score before tailoring)
# ─────────────────────────────────────────────────────────────────────────────
with tab1:
    st.markdown("### 🔍 Job Matcher")
    st.markdown(
        "Paste any job description — from LinkedIn, company websites, or elsewhere. "
        "Score how well your resume matches, then send it straight to **Tailor Resume** with one click."
    )
    st.markdown("---")

    resume_ready_t2 = "user_resume" in st.session_state
    if not resume_ready_t2:
        st.info("👈 Upload your resume in the sidebar first.")

    jm_jd = st.text_area(
        "📋 Paste the job description here",
        height=280,
        placeholder="Paste the full job description…",
        key="jm_jd",
        disabled=not resume_ready_t2,
    )

    # Model selector
    openai_key = os.environ.get("OPENAI_API_KEY", "")
    model_choices = {"Claude Haiku": ("claude", bool(api_key))}
    if _OPENAI_AVAILABLE and openai_key:
        model_choices["OpenAI GPT-4o mini"] = ("openai", True)
    elif not _OPENAI_AVAILABLE:
        model_choices["OpenAI GPT-4o mini (install openai package)"] = ("openai", False)
    else:
        model_choices["OpenAI GPT-4o mini (set OPENAI_API_KEY)"] = ("openai", False)

    selected_label = st.radio(
        "🤖 Score with:",
        list(model_choices.keys()),
        horizontal=True,
        key="jm_model",
        disabled=not resume_ready_t2,
    )
    selected_provider, model_ready = model_choices[selected_label]

    jm_btn = st.button(
        "🔍 Match Against My Resume",
        type="primary",
        disabled=(not model_ready or not resume_ready_t2),
        key="jm_btn",
    )

    if jm_btn:
        if not jm_jd.strip():
            st.error("Please paste a job description first.")
        else:
            base_resume = st.session_state["user_resume"]
            with st.spinner(f"Analysing match ({selected_label})…"):
                try:
                    if selected_provider == "openai":
                        result = score_pasted_jd_openai(base_resume, jm_jd.strip())
                    else:
                        result = score_pasted_jd(base_resume, jm_jd.strip())
                except json.JSONDecodeError as exc:
                    st.error(f"Scoring returned invalid JSON: {exc}")
                    st.stop()
                except Exception as exc:
                    st.error(f"Scoring failed: {exc}")
                    st.stop()

            score = int(result.get("score", 0))
            recommendation = result.get("recommendation", "Skip")

            score_colour = (
                "🟢" if score >= 85 else
                "🟡" if score >= 70 else
                "🟠" if score >= 65 else
                "🔴"
            )
            apply_banner = (
                "✅ **Yes — apply for this role.**"
                if recommendation == "Apply"
                else "❌ **Skip this one.**"
            )

            st.markdown("---")
            res_col1, res_col2 = st.columns([1, 2])
            with res_col1:
                st.metric("Match Score", f"{score_colour} {score}%")
            with res_col2:
                st.markdown(f"### {apply_banner}")
                if result.get("recommendation_reason"):
                    st.markdown(result["recommendation_reason"])

            st.markdown("---")
            if result.get("fit"):
                st.info(f"**Overall fit:** {result['fit']}")
            if result.get("match"):
                st.markdown("**Why you're a good fit:**")
                for m in result["match"]:
                    st.markdown(f"  ✅ {m}")
            if result.get("gap"):
                st.markdown("**Gaps / things to address:**")
                for g in result["gap"]:
                    st.markdown(f"  ⚠️ {g}")
            if result.get("missing_keywords"):
                st.markdown("**ATS keywords missing from your resume:**")
                st.markdown("  " + "  ·  ".join(result["missing_keywords"]))

            # Auto-transfer JD to Tailor Resume tab
            st.session_state["prefill_jd"] = jm_jd.strip()
            st.info("✏️ Job description sent to **Tailor Resume** tab — your details are ready to go!")

            # Save to session profile
            _pd = get_profile_data()
            _pd["analyzed_jds"].append({
                "id": str(uuid.uuid4())[:8],
                "timestamp": datetime.now(timezone.utc).isoformat(),
                "jd_snippet": jm_jd.strip()[:500],
                "score": score,
                "recommendation": recommendation,
                "fit": result.get("fit", ""),
                "match": result.get("match", []),
                "gap": result.get("gap", []),
                "missing_keywords": result.get("missing_keywords", []),
                "recommendation_reason": result.get("recommendation_reason", ""),
            })


# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — Tailor Resume
# ─────────────────────────────────────────────────────────────────────────────
with tab2:
    st.markdown("---")

    resume_ready = "user_resume" in st.session_state

    if not resume_ready:
        st.info("👈 Upload your resume in the sidebar to get started.")

    left, right = st.columns([3, 2])

    with left:
        company_name = st.text_input(
            "🏢 Company Name",
            placeholder="e.g., Google, GovTech, Shopee",
            disabled=not resume_ready,
        )
        role_name = st.text_input(
            "💼 Role / Job Title",
            placeholder="e.g., Product Manager, Senior Programme Manager",
            disabled=not resume_ready,
        )
        # Pre-fill JD if transferred from Job Matcher
        _prefill_jd = st.session_state.get("prefill_jd", "")
        if _prefill_jd:
            st.caption("✅ Job description transferred from Job Matcher.")
        job_description = st.text_area(
            "📋 Paste Job Description",
            value=_prefill_jd,
            height=380,
            placeholder="Paste the full job description here…",
            disabled=not resume_ready,
        )

    with right:
        st.markdown("### What this tool does")
        st.markdown(
            """
**Style applied (value-driven method):**

✅ **6-second rule** — Summary answers *"so what?"* instantly
✅ **Bold lead-ins** — Every bullet opens with a bold metric or achievement
✅ **Metrics front** — Numbers and impact are the first thing recruiters see
✅ **Key Skills** — 6 keywords pulled directly from the JD for ATS
✅ **Anti-fabrication** — Only real facts from your resume, never invented
✅ **ATS-optimised** — Standard headers, keyword-rich, no complex formatting
✅ **New file** — Your original resume is never overwritten
            """
        )

        st.markdown("### Output format")
        st.markdown(
            "Reverse-chronological .docx — Calibri 11pt, section headers "
            "with underline borders, professional bullet layout."
        )
        st.markdown(
            "File: `Resume_{Name}_{Company}_{Date}.docx`"
        )

        also_cover_letter = st.checkbox(
            "✉️ Also generate a Cover Letter",
            value=False,
            disabled=not resume_ready,
        )

        _MAX_RUNS = 3
        if "tailor_runs_used" not in st.session_state:
            st.session_state["tailor_runs_used"] = 0
        _runs_used = st.session_state["tailor_runs_used"]
        _runs_left = _MAX_RUNS - _runs_used

        if _runs_left > 0:
            st.caption(f"Tailoring runs remaining this session: **{_runs_left} / {_MAX_RUNS}**")
        else:
            st.warning(f"You have used all {_MAX_RUNS} tailoring runs for this session. Please refresh the page to start a new session.")

        run_btn = st.button(
            "🚀 Tailor My Resume",
            type="primary",
            use_container_width=True,
            disabled=(not api_key or not resume_ready or _runs_left <= 0),
        )

    st.markdown("---")

    # ── Run tailoring ─────────────────────────────────────────────────────────
    if run_btn:
        errors = []
        if not company_name.strip():
            errors.append("Company name is required.")
        if not role_name.strip():
            errors.append("Role / Job Title is required.")
        if not job_description.strip():
            errors.append("Job description cannot be empty.")

        if errors:
            for e in errors:
                st.error(e)
        else:
            st.session_state.pop("tailor_result", None)
            base_resume = st.session_state["user_resume"]

            try:
                _openai_refine = _OPENAI_AVAILABLE and bool(os.environ.get("OPENAI_API_KEY"))
                _refine_steps  = 1 if _openai_refine else 0
                total_steps    = (4 if also_cover_letter else 2) + _refine_steps
                date_str       = datetime.now().strftime("%Y%m%d")
                safe_company   = re.sub(r'[\\/*?:"<>|]', "", company_name).strip()
                safe_role      = re.sub(r'[\\/*?:"<>|]', "", role_name).strip()

                # Pre-step: gap analysis
                _jd_gaps: list = []
                _jd_missing_kw: list = []
                with st.spinner("Analysing JD gaps to inform tailoring…"):
                    try:
                        _match = score_pasted_jd(base_resume, job_description)
                        _jd_gaps = _match.get("gap", [])
                        _jd_missing_kw = _match.get("missing_keywords", [])
                        if _jd_gaps or _jd_missing_kw:
                            st.caption(
                                f"Gap analysis complete — {len(_jd_gaps)} gap(s) and "
                                f"{len(_jd_missing_kw)} missing keyword(s) identified. "
                                "Claude will address these in the tailored output."
                            )
                    except Exception:
                        pass  # Non-fatal

                # Step 1: Claude tailors
                st.session_state["tailor_runs_used"] = st.session_state.get("tailor_runs_used", 0) + 1
                st.info(f"Step 1 / {total_steps}  Claude is tailoring your resume…")
                tailored = tailor_with_claude(
                    base_resume, job_description, company_name, role_name,
                    gaps=_jd_gaps, missing_keywords=_jd_missing_kw,
                )

                # Save to session profile
                _pd = get_profile_data()
                _pd["tailored_roles"].append({
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                    "company": company_name,
                    "role": role_name,
                    "jd_snippet": job_description[:500],
                })

                # Step 2 (optional): OpenAI refinement
                _next_step = 2
                if _openai_refine:
                    st.info(f"Step {_next_step} / {total_steps}  OpenAI GPT-4o is reviewing Claude's output…")
                    try:
                        tailored = refine_with_openai(
                            base_resume, tailored, job_description,
                            gaps=_jd_gaps, missing_keywords=_jd_missing_kw,
                        )
                        st.caption("✨ OpenAI refinement complete.")
                    except Exception as _oe:
                        st.warning(f"OpenAI refinement skipped ({_oe}). Using Claude's version.")
                    _next_step += 1

                # Build resume docx
                st.info(f"Step {_next_step} / {total_steps}  Building your formatted .docx…")
                docx_bytes = build_docx(tailored)
                candidate_name = tailored.get("name", "Candidate")
                safe_name = re.sub(r'[\\/*?:"<>|]', "", candidate_name).strip()
                resume_filename = f"Resume_{safe_name}_{safe_company}_{date_str}.docx"
                _next_step += 1

                # Cover letter (optional)
                cl_bytes, cl_filename, cl_data = None, None, None
                if also_cover_letter:
                    st.info(f"Step {_next_step} / {total_steps}  Claude is writing your cover letter…")
                    cl_data = generate_cover_letter(
                        base_resume, job_description, company_name, role_name,
                        gaps=_jd_gaps, missing_keywords=_jd_missing_kw,
                    )
                    _next_step += 1
                    st.info(f"Step {_next_step} / {total_steps}  Building cover letter .docx…")
                    cl_bytes = build_cover_letter_docx(cl_data, company_name, role_name)
                    cl_filename = f"Cover Letter_{safe_name}_{safe_company}_{date_str}.docx"

                st.session_state["tailor_result"] = {
                    "docx_bytes": docx_bytes,
                    "resume_filename": resume_filename,
                    "tailored": tailored,
                    "cl_bytes": cl_bytes,
                    "cl_filename": cl_filename,
                    "cl_data": cl_data,
                    "also_cover_letter": also_cover_letter,
                }

            except json.JSONDecodeError as exc:
                st.error(f"Claude returned invalid JSON: {exc}")
                st.code(traceback.format_exc())
            except anthropic.AuthenticationError:
                st.error("Invalid API key. Check your ANTHROPIC_API_KEY.")
            except Exception as exc:
                st.error(f"Unexpected error: {exc}")
                st.code(traceback.format_exc())

    # ── Download & Preview ────────────────────────────────────────────────────
    if "tailor_result" in st.session_state:
        _res = st.session_state["tailor_result"]
        _tailored = _res["tailored"]

        st.success("✅ Resume tailored and ready to download!")
        st.download_button(
            label=f"⬇️  Download  {_res['resume_filename']}",
            data=_res["docx_bytes"],
            file_name=_res["resume_filename"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="resume_download",
        )

        if _res.get("also_cover_letter") and _res.get("cl_bytes"):
            st.success("✅ Cover letter ready!")
            st.download_button(
                label=f"⬇️  Download  {_res['cl_filename']}",
                data=_res["cl_bytes"],
                file_name=_res["cl_filename"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="cl_download",
            )

        st.markdown("---")
        st.markdown("### Resume Preview")

        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**Title Line**")
            st.info(_tailored.get("title_line", "-"))
            st.markdown("**Professional Summary**")
            st.info(_tailored.get("summary", "-"))

        with c2:
            st.markdown("**Key Skills**")
            for skill in _tailored.get("key_skills", []):
                st.markdown(f"  ✦ {skill}")

        # Cover letter preview
        _cl_data = _res.get("cl_data")
        if _res.get("also_cover_letter") and _cl_data:
            st.markdown("---")
            st.markdown("### Cover Letter Preview")
            st.markdown(f"**{_cl_data.get('subject_line', '')}**")
            st.markdown("Dear Hiring Manager,")
            for para in _cl_data.get("paragraphs", []):
                st.markdown(para)
            name_out = _cl_data.get("signoff_name", "")
            if name_out:
                st.markdown(f"*Sincerely, {name_out}*")




# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — Analysis
# ─────────────────────────────────────────────────────────────────────────────
with tab3:
    st.markdown("---")
    st.markdown(
        "This tab learns from every job description you analyse and every resume you tailor "
        "during this session. The more you use the tool, the sharper the insights."
    )

    resume_ready_t3 = "user_resume" in st.session_state
    if not resume_ready_t3:
        st.info("👈 Upload your resume in the sidebar first.")
    else:
        _profile = get_profile_data()
        _analyzed = _profile.get("analyzed_jds", [])
        _tailored_list = _profile.get("tailored_roles", [])

        if _analyzed:
            _avg_score  = sum(j["score"] for j in _analyzed) / len(_analyzed)
            _apply_count = sum(1 for j in _analyzed if j["recommendation"] == "Apply")
            _skip_count  = len(_analyzed) - _apply_count

            col_a, col_b, col_c, col_d = st.columns(4)
            col_a.metric("JDs Analysed",    len(_analyzed))
            col_b.metric("Avg Match Score", f"{_avg_score:.0f}%")
            col_c.metric("Apply",           _apply_count)
            col_d.metric("Skip",            _skip_count)

            if _tailored_list:
                st.caption(f"Resumes tailored this session: {len(_tailored_list)}")

            st.markdown("---")

        # Profile analysis display
        _analysis = _profile.get("profile_analysis", {})
        if _analysis.get("analysis_text"):
            st.markdown(_analysis["analysis_text"])
            st.caption(f"Last updated: {_analysis.get('last_updated', 'N/A')[:10]}")
        elif not _analyzed:
            st.info(
                "No data yet — analyse some job descriptions in the **Job Matcher** tab first. "
                "Each analysis is saved for the duration of this session."
            )

        if _analyzed:
            st.markdown("---")
            if st.button("🔄 Refresh Analysis", use_container_width=True, key="refresh_analysis"):
                with st.spinner("Generating your career profile analysis…"):
                    base_resume = st.session_state["user_resume"]
                    _analysis_text = generate_profile_analysis(_profile, base_resume)
                    _profile["profile_analysis"] = {
                        "last_updated": datetime.now(timezone.utc).isoformat(),
                        "analysis_text": _analysis_text,
                    }
                st.rerun()

            st.markdown("---")
            st.markdown("### Session History")
            st.caption("Job descriptions analysed this session (newest first)")
            for _entry in reversed(_analyzed[-20:]):
                _s   = _entry["score"]
                _c   = "🟢" if _s >= 85 else "🟡" if _s >= 70 else "🟠" if _s >= 65 else "🔴"
                _rec = "✅ Apply" if _entry["recommendation"] == "Apply" else "❌ Skip"
                _date = _entry["timestamp"][:10]
                _snip = _entry.get("jd_snippet", "")[:100].replace("\n", " ")
                st.markdown(
                    f"{_c} **{_s}%** &nbsp;|&nbsp; {_rec} &nbsp;|&nbsp; "
                    f"{_date} &nbsp;|&nbsp; _{_snip}…_"
                )
